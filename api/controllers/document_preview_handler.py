from typing import Dict, Any, Optional
import html
import re
from pathlib import Path
from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.table import Table

class DocumentPreviewHandler:
    """CKDEV-NOTE: Handler for generating clean HTML previews from DOCX templates"""
    
    def __init__(self, template_path: str, replacements: Dict[str, str] = None):
        self.template_path = Path(template_path)
        self.replacements = replacements or {}
        self.doc = DocxDocument(self.template_path)
        
    def escape_html(self, text: str) -> str:
        return html.escape(text or '')
    
    def highlight_placeholders(self, raw_text: str) -> str:
        """Enhanced placeholder highlighting with badge-style components"""
        if not raw_text:
            return ''
        text = self.escape_html(raw_text)
        if not self.replacements:
            return text
        
        # CKDEV-NOTE: Replace longer keys first to avoid partial overlaps, with enhanced badge styling
        keys = sorted(self.replacements.keys(), key=lambda k: len(k), reverse=True)
        for key in keys:
            if not key:
                continue
            value = (self.replacements.get(key) or '').strip()
            safe_key = self.escape_html(key)
            if value:
                safe_val = self.escape_html(value)
                # Filled placeholder with success badge styling
                text = text.replace(safe_key, f'''
                <span class="inline-flex items-center rounded-md border border-emerald-200 bg-emerald-50 px-2 py-0.5 text-xs font-medium text-emerald-700 transition-colors">
                    {safe_val}
                </span>''')
            else:
                # Empty placeholder with warning badge styling
                text = text.replace(safe_key, f'''
                <span class="inline-flex items-center rounded-md border border-amber-200 bg-amber-50 px-2 py-0.5 text-xs font-medium text-amber-700 animate-pulse">
                    {safe_key}
                </span>''')
        return text
    
    def iter_block_items(self, parent):
        """Iterate through all blocks (paragraphs and tables) in order"""
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        
        for child in parent.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)
    
    def render_table_as_grid(self, table) -> str:
        """Render table as structured grid instead of HTML table"""
        html_parts = []
        
        # Process each row as structured content
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                cell_text = " ".join([p.text for p in cell.paragraphs]).strip()
                if cell_text:
                    row_data.append(cell_text)
            
            if row_data:
                # Check if this looks like a label-value pair
                if len(row_data) == 2:
                    label, value = row_data
                    # Skip if both are same (duplicated content)
                    if label != value:
                        html_parts.append(f'''
                        <div class="info-row">
                            <span class="info-label">{self.highlight_placeholders(label)}</span>
                            <span class="info-value">{self.highlight_placeholders(value)}</span>
                        </div>''')
                elif len(row_data) == 1:
                    # Single cell, treat as paragraph or section header
                    text = row_data[0]
                    if text.isupper() and len(text) < 50:
                        html_parts.append(f'<h2 class="section-title">{self.highlight_placeholders(text)}</h2>')
                    else:
                        html_parts.append(f'<p class="declaration-text">{self.highlight_placeholders(text)}</p>')
                else:
                    # Multiple cells, group as a section
                    for item in row_data:
                        html_parts.append(f'<p class="declaration-text">{self.highlight_placeholders(item)}</p>')
        
        if html_parts:
            return f'<div class="info-section">{"".join(html_parts)}</div>'
        return ''
    
    def extract_table_content(self, table):
        """Extract meaningful content from table with enhanced styling and justification"""
        content_parts = []
        seen_content = set()
        
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in seen_content:
                    row_cells.append(cell_text)
                    seen_content.add(cell_text)
            
            # Process meaningful row content with enhanced styling
            if row_cells:
                # Check for label-value pairs
                if len(row_cells) >= 2:
                    for i in range(0, len(row_cells), 2):
                        if i + 1 < len(row_cells):
                            label = row_cells[i]
                            value = row_cells[i + 1]
                            if label != value and not label.startswith('{{') and not value.startswith('{{'):
                                content_parts.append(f'''
                                <div class="flex flex-wrap gap-2 items-center mb-3">
                                    <span class="inline-flex items-center rounded-md border border-slate-200 bg-slate-50 px-2.5 py-0.5 text-xs font-medium text-slate-700">
                                        {self.highlight_placeholders(label)}
                                    </span>
                                    <span class="text-sm font-medium text-slate-900">
                                        {self.highlight_placeholders(value)}
                                    </span>
                                </div>''')
                            elif '{{' in label or '{{' in value:
                                content_parts.append(f'<p class="text-justify my-3 leading-relaxed hyphens-auto">{self.highlight_placeholders(label)} {self.highlight_placeholders(value)}</p>')
                else:
                    # Single content with better typography
                    content_parts.append(f'<p class="text-justify my-3 leading-relaxed hyphens-auto indent-4">{self.highlight_placeholders(row_cells[0])}</p>')
        
        return content_parts

    def generate_structured_html(self) -> str:
        """Generate clean, structured HTML with enhanced typography and Card-style layout"""
        html_parts = []
        
        # CKDEV-NOTE: Enhanced Card-style layout with proper justification and typography
        html_parts.append('''
        <div class="max-w-[800px] mx-auto">
            <div class="bg-white text-slate-900 flex flex-col gap-6 rounded-xl border border-slate-200 py-6 shadow-sm print:shadow-none print:border-none print:p-0 
                       font-['Times_New_Roman',serif]">
        ''')
        
        # Document Header with Card styling
        html_parts.append('''
            <div class="px-6 border-b border-slate-100 pb-6 print:border-none print:pb-0">
        ''')
        
        # CKDEV-NOTE: Extract document title paragraphs and separate location/date info
        location_date_elements = []
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Check if this is a location/date element (contains city name or date pattern)
                if any(city in text.upper() for city in ['SÃO JOSÉ DOS CAMPOS', 'SÃO PAULO', 'CAMPINAS']) or ('{{' in text and ('DATA' in text.upper() or 'LOCAL' in text.upper())):
                    # Store for rendering at the end
                    location_date_elements.append(text)
                else:
                    # Regular title - render normally
                    html_parts.append(f'''
                    <h1 class="text-xl font-semibold text-center mb-4 tracking-wide uppercase text-slate-800 leading-tight">
                        {self.highlight_placeholders(text)}
                    </h1>''')
        
        html_parts.append('</div>')
        
        # Document Content with Card content styling
        html_parts.append('''
            <div class="px-6">
        ''')
        
        # Extract meaningful content from tables with enhanced sections
        for i, table in enumerate(self.doc.tables):
            table_content = self.extract_table_content(table)
            
            # Add structured table content
            unique_content = []
            seen_lines = set()
            
            for content in table_content:
                # Remove HTML tags for comparison
                import re
                clean_content = re.sub('<[^<]+?>', '', content).strip()
                if clean_content and clean_content not in seen_lines:
                    unique_content.append(content)
                    seen_lines.add(clean_content)
            
            # Add organized sections with Card-style grouping
            if unique_content:
                # Section separator
                if i > 0:
                    html_parts.append('<div class="h-px w-full bg-slate-200 my-6"></div>')
                
                # Content section with enhanced styling
                html_parts.append(f'''
                <div class="bg-slate-50 rounded-lg border border-slate-100 p-6 my-6">
                    <div class="space-y-4">''')
                
                for content in unique_content:
                    html_parts.append(content)
                
                html_parts.append('''
                    </div>
                </div>''')
        
        html_parts.append('</div>')
        
        # CKDEV-NOTE: Add location/date elements before signature section
        if location_date_elements:
            html_parts.append('''
                <div class="px-6 py-4 border-t border-slate-100">''')
            for location_date_text in location_date_elements:
                html_parts.append(f'''
                    <h1 class="text-xl font-semibold text-center mb-4 tracking-wide uppercase text-slate-800 leading-tight">
                        {self.highlight_placeholders(location_date_text)}
                    </h1>''')
            html_parts.append('</div>')
        
        # Enhanced signature section with Card footer styling
        html_parts.append('''
            <div class="flex items-center justify-center px-6 border-t border-slate-100 pt-6 print:border-none print:pt-8">
                <div class="text-center space-y-3">
                    <div class="w-64 h-px bg-slate-300 mx-auto"></div>
                    <div class="space-y-1">
                        <p class="text-xs font-medium text-slate-600">Assinatura</p>
                        <p class="text-xs text-slate-500">Nome: ___________________________</p>
                        <p class="text-xs text-slate-500">CPF: ___________________________</p>
                    </div>
                </div>
            </div>
        ''')
        
        html_parts.append('''
            </div>
        </div>''')
        
        return '\n'.join(html_parts)