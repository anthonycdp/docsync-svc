import os
import platform
import tempfile
from pathlib import Path
from typing import Tuple, Optional
import logging

# CKDEV-NOTE: Enhanced DOCX to PDF converter preserving formatting and layout
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    from weasyprint import HTML, CSS
    from weasyprint.text.fonts import FontConfiguration
except ImportError:
    HTML = None

try:
    from docx2pdf import convert as docx2pdf_convert
except ImportError:
    docx2pdf_convert = None

try:
    from .libreoffice_converter import convert_docx_to_pdf_libreoffice
except ImportError:
    try:
        from libreoffice_converter import convert_docx_to_pdf_libreoffice
    except ImportError:
        convert_docx_to_pdf_libreoffice = None


class EnhancedDocxToPdfConverter:
    """
    Enhanced DOCX to PDF converter that preserves original formatting and layout.
    Uses multiple conversion methods with priority for layout preservation.
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self._timeout = 60
    
    def convert_docx_to_pdf(
        self, 
        docx_path: str, 
        pdf_path: Optional[str] = None
    ) -> Tuple[bool, str, Optional[str]]:
        """
        Convert DOCX to PDF preserving original formatting and layout.
        
        Priority order:
        1. LibreOffice (best formatting preservation)
        2. docx2pdf (Windows, good formatting)
        3. DOCX -> HTML -> PDF (custom formatting preservation)
        4. Basic python-docx fallback
        
        Args:
            docx_path: Path to the .docx file
            pdf_path: Output path for PDF (optional)
        
        Returns:
            Tuple[bool, str, Optional[str]]: (success, message, pdf_path)
        """
        if not os.path.exists(docx_path):
            return False, f"DOCX file not found: {docx_path}", None
        
        if pdf_path is None:
            pdf_path = str(Path(docx_path).with_suffix('.pdf'))
        
        # CKDEV-NOTE: Method 1 - LibreOffice (best layout preservation)
        if convert_docx_to_pdf_libreoffice:
            try:
                success, message, result_path = convert_docx_to_pdf_libreoffice(docx_path, pdf_path)
                if success and result_path and self._validate_pdf_quality(result_path):
                    self.logger.info(f"LibreOffice conversion successful with good layout: {result_path}")
                    return success, f"LibreOffice (high fidelity): {message}", result_path
                else:
                    self.logger.warning(f"LibreOffice conversion failed or poor quality: {message}")
            except Exception as e:
                self.logger.warning(f"LibreOffice conversion error: {e}")
        
        # CKDEV-NOTE: Method 2 - docx2pdf (Windows, maintains good formatting)
        if docx2pdf_convert and platform.system().lower() == 'windows':
            try:
                success, message, result_path = self._convert_with_docx2pdf(docx_path, pdf_path)
                if success and result_path and self._validate_pdf_quality(result_path):
                    self.logger.info(f"docx2pdf conversion successful: {result_path}")
                    return success, f"docx2pdf (Windows): {message}", result_path
                else:
                    self.logger.warning(f"docx2pdf conversion failed or poor quality: {message}")
            except Exception as e:
                self.logger.warning(f"docx2pdf conversion error: {e}")
        
        # CKDEV-NOTE: Method 3 - DOCX to HTML to PDF (custom formatting preservation)
        if Document and HTML:
            try:
                success, message, result_path = self._convert_via_html(docx_path, pdf_path)
                if success and result_path:
                    self.logger.info(f"HTML-based conversion successful: {result_path}")
                    return success, f"HTML-based (custom): {message}", result_path
                else:
                    self.logger.warning(f"HTML-based conversion failed: {message}")
            except Exception as e:
                self.logger.warning(f"HTML-based conversion error: {e}")
        
        # CKDEV-NOTE: Method 4 - Enhanced python-docx fallback
        if Document:
            try:
                success, message, result_path = self._convert_with_enhanced_formatting(docx_path, pdf_path)
                if success and result_path:
                    self.logger.info(f"Enhanced python-docx conversion successful: {result_path}")
                    return success, f"Enhanced python-docx: {message}", result_path
                else:
                    self.logger.warning(f"Enhanced python-docx conversion failed: {message}")
            except Exception as e:
                self.logger.warning(f"Enhanced python-docx conversion error: {e}")
        
        return False, "All PDF conversion methods failed to preserve proper formatting", None
    
    def _validate_pdf_quality(self, pdf_path: str) -> bool:
        """Validate PDF was generated with reasonable quality"""
        try:
            if not os.path.exists(pdf_path):
                return False
            
            file_size = os.path.getsize(pdf_path)
            # PDFs should be at least 1KB for document content
            if file_size < 1024:
                return False
            
            # Basic PDF header check
            with open(pdf_path, 'rb') as f:
                header = f.read(8)
                if not header.startswith(b'%PDF'):
                    return False
            
            return True
        except Exception:
            return False
    
    def _convert_with_docx2pdf(self, docx_path: str, pdf_path: str) -> Tuple[bool, str, Optional[str]]:
        """Convert using docx2pdf with enhanced error handling"""
        if not docx2pdf_convert:
            return False, "docx2pdf library not available", None
        
        try:
            # CKDEV-NOTE: Clean up existing PDF
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
            
            # CKDEV-NOTE: Ensure output directory exists
            output_dir = os.path.dirname(pdf_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
            
            # CKDEV-NOTE: Use absolute paths
            docx_path_abs = os.path.abspath(docx_path)
            pdf_path_abs = os.path.abspath(pdf_path)
            
            self.logger.info(f"Converting with docx2pdf: {docx_path_abs}")
            
            # CKDEV-NOTE: Perform conversion with timeout
            import threading
            
            conversion_result = {'success': False, 'error': None}
            
            def convert_worker():
                try:
                    docx2pdf_convert(docx_path_abs, pdf_path_abs)
                    conversion_result['success'] = True
                except Exception as e:
                    conversion_result['error'] = e
            
            thread = threading.Thread(target=convert_worker)
            thread.daemon = True
            thread.start()
            thread.join(timeout=self._timeout)
            
            if thread.is_alive():
                return False, f"docx2pdf conversion timed out after {self._timeout} seconds", None
            
            if conversion_result.get('error'):
                raise conversion_result['error']
            
            # CKDEV-NOTE: Validate generated PDF
            if not os.path.exists(pdf_path_abs) or os.path.getsize(pdf_path_abs) == 0:
                return False, "docx2pdf generated empty or no PDF file", None
            
            file_size = os.path.getsize(pdf_path_abs)
            return True, f"PDF generated with good formatting ({file_size} bytes)", pdf_path_abs
            
        except Exception as e:
            self.logger.error(f"docx2pdf conversion error: {e}")
            return False, f"docx2pdf conversion error: {str(e)}", None
    
    def _convert_via_html(self, docx_path: str, pdf_path: str) -> Tuple[bool, str, Optional[str]]:
        """Convert DOCX to PDF via HTML for better formatting preservation"""
        if not Document or not HTML:
            return False, "Required libraries (python-docx, weasyprint) not available", None
        
        try:
            # CKDEV-NOTE: Parse DOCX document
            doc = Document(docx_path)
            
            # CKDEV-NOTE: Convert to HTML with formatting preservation
            html_content = self._docx_to_html(doc)
            
            # CKDEV-NOTE: Create CSS for proper formatting
            css_styles = self._generate_css_styles()
            
            # CKDEV-NOTE: Generate PDF from HTML
            font_config = FontConfiguration()
            html_doc = HTML(string=html_content)
            css_doc = CSS(string=css_styles, font_config=font_config)
            
            html_doc.write_pdf(pdf_path, stylesheets=[css_doc], font_config=font_config)
            
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                file_size = os.path.getsize(pdf_path)
                return True, f"HTML-based conversion with preserved formatting ({file_size} bytes)", pdf_path
            else:
                return False, "HTML-based conversion generated empty PDF", None
                
        except Exception as e:
            self.logger.error(f"HTML-based conversion error: {e}")
            return False, f"HTML-based conversion error: {str(e)}", None
    
    def _docx_to_html(self, doc) -> str:
        """Convert DOCX document to HTML preserving structure and formatting"""
        html_parts = ['<!DOCTYPE html><html><head><meta charset="utf-8"><title>Document</title></head><body>']
        
        for paragraph in doc.paragraphs:
            # CKDEV-NOTE: Determine paragraph style and alignment
            style_class = self._get_paragraph_style_class(paragraph)
            alignment = self._get_paragraph_alignment(paragraph)
            
            # CKDEV-NOTE: Handle different paragraph types
            if paragraph.style.name.startswith('Heading') or self._is_title_paragraph(paragraph):
                level = self._extract_heading_level(paragraph.style.name)
                html_parts.append(f'<h{level} class="{style_class}" style="text-align: {alignment};">')
                html_parts.append(self._process_paragraph_runs(paragraph))
                html_parts.append(f'</h{level}>')
            else:
                html_parts.append(f'<p class="{style_class}" style="text-align: {alignment};">')
                html_parts.append(self._process_paragraph_runs(paragraph))
                html_parts.append('</p>')
        
        html_parts.append('</body></html>')
        return ''.join(html_parts)
    
    def _get_paragraph_style_class(self, paragraph) -> str:
        """Get CSS class for paragraph based on its style"""
        style_name = paragraph.style.name.lower().replace(' ', '-')
        return f"style-{style_name}"
    
    def _get_paragraph_alignment(self, paragraph) -> str:
        """Get text alignment for paragraph"""
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return 'center'
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return 'right'
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return 'justify'
        else:
            return 'left'
    
    def _is_title_paragraph(self, paragraph) -> bool:
        """Check if paragraph should be treated as title"""
        text = paragraph.text.upper()
        return 'TERMO DE RESPONSABILIDADE' in text or text.startswith('TERMO')
    
    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name"""
        if 'heading 1' in style_name.lower() or self._is_title_paragraph:
            return 1
        elif 'heading 2' in style_name.lower():
            return 2
        elif 'heading 3' in style_name.lower():
            return 3
        else:
            return 1
    
    def _process_paragraph_runs(self, paragraph) -> str:
        """Process paragraph runs to preserve formatting"""
        html_parts = []
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # CKDEV-NOTE: Apply run-level formatting
            if run.bold and run.italic:
                text = f'<strong><em>{text}</em></strong>'
            elif run.bold:
                text = f'<strong>{text}</strong>'
            elif run.italic:
                text = f'<em>{text}</em>'
            
            if run.underline:
                text = f'<u>{text}</u>'
            
            html_parts.append(text)
        
        return ''.join(html_parts)
    
    def _generate_css_styles(self) -> str:
        """Generate CSS styles for proper document formatting"""
        return """
        @page {
            size: A4;
            margin: 2cm;
        }
        
        body {
            font-family: 'Times New Roman', serif;
            font-size: 12pt;
            line-height: 1.6;
            color: black;
            margin: 0;
            padding: 0;
        }
        
        h1 {
            font-size: 14pt;
            font-weight: bold;
            text-align: center;
            margin: 0 0 20pt 0;
            padding: 0;
        }
        
        p {
            margin: 0 0 12pt 0;
            text-align: justify;
            text-indent: 0;
        }
        
        .style-title {
            font-size: 14pt;
            font-weight: bold;
            text-align: center;
            margin: 0 0 20pt 0;
        }
        
        .style-normal {
            font-size: 12pt;
            margin: 0 0 12pt 0;
        }
        
        strong {
            font-weight: bold;
        }
        
        em {
            font-style: italic;
        }
        
        u {
            text-decoration: underline;
        }
        """
    
    def _convert_with_enhanced_formatting(self, docx_path: str, pdf_path: str) -> Tuple[bool, str, Optional[str]]:
        """Enhanced python-docx conversion with better formatting preservation"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import inch, cm
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
            
            # CKDEV-NOTE: Load DOCX document
            doc = Document(docx_path)
            
            # CKDEV-NOTE: Create PDF document
            pdf_doc = SimpleDocTemplate(
                pdf_path,
                pagesize=A4,
                rightMargin=2*cm,
                leftMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )
            
            # CKDEV-NOTE: Create styles
            styles = getSampleStyleSheet()
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=14,
                spaceAfter=20,
                alignment=TA_CENTER,
                fontName='Times-Bold'
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12,
                alignment=TA_JUSTIFY,
                fontName='Times-Roman',
                leading=18
            )
            
            # CKDEV-NOTE: Build content
            story = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # CKDEV-NOTE: Determine if this is a title
                if i == 0 or 'TERMO DE RESPONSABILIDADE' in text.upper():
                    story.append(Paragraph(text, title_style))
                else:
                    story.append(Paragraph(text, normal_style))
                    
                # CKDEV-NOTE: Add spacing between paragraphs
                if i < len(doc.paragraphs) - 1:
                    story.append(Spacer(1, 6))
            
            # CKDEV-NOTE: Build PDF
            pdf_doc.build(story)
            
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                file_size = os.path.getsize(pdf_path)
                return True, f"Enhanced formatting conversion successful ({file_size} bytes)", pdf_path
            else:
                return False, "Enhanced formatting conversion generated empty PDF", None
                
        except ImportError as e:
            return False, f"Missing library for enhanced conversion: {str(e)}", None
        except Exception as e:
            self.logger.error(f"Enhanced formatting conversion error: {e}")
            return False, f"Enhanced formatting conversion error: {str(e)}", None


def convert_docx_to_pdf_enhanced(
    docx_path: str, 
    pdf_path: Optional[str] = None
) -> Tuple[bool, str, Optional[str]]:
    """
    Enhanced DOCX to PDF conversion preserving formatting and layout.
    
    Args:
        docx_path: Path to the .docx file
        pdf_path: Output path for PDF (optional)
    
    Returns:
        Tuple[bool, str, Optional[str]]: (success, message, pdf_path)
    """
    converter = EnhancedDocxToPdfConverter()
    return converter.convert_docx_to_pdf(docx_path, pdf_path)