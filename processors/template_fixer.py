import os
import re
from docx import Document
from docx.shared import Pt
from typing import Dict, Tuple

class TemplateFixer:
    
    def __init__(self, placeholder_fixes: Dict[str, str] = None):
        default_fixes = {
            r'ANO_(?!MODELO_VEICULO)[A-Z0-9\s\.]+': 'ANO_MODELO_VEICULO',
            r'NOME_TRACKER': 'MODELO_VEICULO',
            r'COR_TRACKER': 'COR_VEICULO',
        }
        
        env_fixes = os.getenv('TEMPLATE_FIXER_PLACEHOLDER_FIXES', '')
        if env_fixes:
            try:
                custom_fixes = {}
                for fix in env_fixes.split(','):
                    if ':' in fix:
                        pattern, replacement = fix.split(':', 1)
                        custom_fixes[pattern.strip()] = replacement.strip()
                self.placeholder_fixes = custom_fixes if custom_fixes else default_fixes
            except Exception:
                self.placeholder_fixes = default_fixes
        else:
            self.placeholder_fixes = placeholder_fixes or default_fixes
    
    def fix_template(self, template_path: str, output_path: str = None):
        doc = Document(template_path)
        fixes_made = 0
        
        for paragraph in doc.paragraphs:
            fixes = self._fix_paragraph(paragraph)
            fixes_made += fixes
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        fixes = self._fix_paragraph(paragraph)
                        fixes_made += fixes
        
        if output_path:
            doc.save(output_path)
        
        return doc, fixes_made
    
    def _fix_paragraph(self, paragraph) -> int:
        if not paragraph.text:
            return 0
        
        full_text = "".join(run.text for run in paragraph.runs)
        original_text = full_text
        fixes_made = 0
        
        for pattern, replacement in self.placeholder_fixes.items():
            matches = re.findall(pattern, full_text)
            for match in matches:
                if match != replacement:
                    full_text = full_text.replace(match, replacement)
                    fixes_made += 1
        
        if full_text != original_text:
            font_name, font_size = self._get_font_settings(paragraph)
            
            while paragraph.runs:
                paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
            
            new_run = paragraph.add_run(full_text)
            new_run.font.name = font_name
            new_run.font.size = font_size
        
        return fixes_made
    
    def _get_font_settings(self, paragraph) -> Tuple[str, Pt]:
        if paragraph.runs:
            font_name = paragraph.runs[0].font.name or os.getenv('DEFAULT_FONT_NAME', 'Aptos')
            font_size = paragraph.runs[0].font.size or Pt(int(os.getenv('DEFAULT_FONT_SIZE', '11')))
        else:
            font_name = os.getenv('DEFAULT_FONT_NAME', 'Aptos')
            font_size = Pt(int(os.getenv('DEFAULT_FONT_SIZE', '11')))
        
        return font_name, font_size

def apply_template_fixes(template_path: str = None, output_path: str = None):
    default_template_path = os.getenv('TEMPLATE_FIXER_INPUT_PATH', '')
    default_output_path = os.getenv('TEMPLATE_FIXER_OUTPUT_PATH', '')
    
    template_path = template_path or default_template_path
    output_path = output_path or default_output_path
    
    if not template_path:
        return None
    
    if not output_path:
        output_path = None
    
    fixer = TemplateFixer()
    doc, fixes = fixer.fix_template(template_path, output_path)
    
    return output_path or template_path

if __name__ == "__main__":
    apply_template_fixes()