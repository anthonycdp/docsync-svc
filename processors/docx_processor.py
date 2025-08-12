"""Módulo para processamento de templates DOCX"""
import os, re
from typing import Dict
from docx import Document
from docx.shared import Pt
from datetime import datetime

try:
    from ..data import ExtractedData
    from ..utils import get_brand_lookup, LoggerMixin
    from .template_replacements import TemplateReplacementManager
except ImportError:
    import sys
    from pathlib import Path
    sys.path.insert(0, str(Path(__file__).parent.parent))
    from data import ExtractedData
    from utils import get_brand_lookup, LoggerMixin
    from processors.template_replacements import TemplateReplacementManager

class DOCXProcessor(LoggerMixin):
    def __init__(self, template_path: str):
        super().__init__()
        self.template_path = template_path
        if not os.path.exists(template_path):
            from ..utils.exceptions import TemplateNotFoundError
            raise TemplateNotFoundError(template_path)
        self.replacement_manager = TemplateReplacementManager()
        self.current_data = None  # CKDEV-NOTE: Store current extracted data for filename generation
    
    def generate_document(self, data: ExtractedData, output_path: str) -> str:
        # CKDEV-NOTE: Store current data for alternative filename generation
        self.current_data = data
        doc = None
        try:
            doc = Document(self.template_path)
            replacements = self._prepare_replacements(data)
            self._replace_text_in_document(doc, replacements)
            
            doc.save(output_path)
            
            try:
                from ..utils.pdf_converter import convert_docx_to_pdf
            except ImportError:
                from utils.pdf_converter import convert_docx_to_pdf
            
            pdf_success, pdf_message, pdf_path = convert_docx_to_pdf(output_path)
            if pdf_success:
                self.logger.info(f"PDF generated successfully: {pdf_path}")
            else:
                self.logger.warning(f"PDF generation failed: {pdf_message}")
            
            return output_path
            
        except FileNotFoundError as e:
            try:
                from ..utils.exceptions import DocumentProcessingError
            except ImportError:
                from utils.exceptions import DocumentProcessingError
            raise DocumentProcessingError(f"Template não acessível: {self.template_path}", template_path=self.template_path) from e
        except PermissionError as e:
            alternative_path = self._generate_alternative_filename(output_path)
            try:
                if doc:
                    doc.save(alternative_path)
                else:
                    doc = Document(self.template_path)
                    replacements = self._prepare_replacements(data)
                    self._replace_text_in_document(doc, replacements)
                    doc.save(alternative_path)
                
                try:
                    from ..utils import convert_docx_to_pdf_libreoffice
                except ImportError:
                    from utils import convert_docx_to_pdf_libreoffice
                
                pdf_success, pdf_message, pdf_path = convert_docx_to_pdf_libreoffice(alternative_path)
                if pdf_success:
                    self.logger.info(f"PDF generated successfully: {pdf_path}")
                else:
                    self.logger.warning(f"PDF generation failed: {pdf_message}")
                
                return alternative_path
            except Exception:
                try:
                    from ..utils.exceptions import DocumentProcessingError
                except ImportError:
                    from utils.exceptions import DocumentProcessingError
                raise DocumentProcessingError(f"Arquivo bloqueado (provavelmente aberto no Word): {output_path}. Feche o arquivo e tente novamente.", template_path=self.template_path) from e
        except Exception as e:
            try:
                from ..utils.exceptions import DocumentProcessingError
            except ImportError:
                from utils.exceptions import DocumentProcessingError
            raise DocumentProcessingError(f"Erro inesperado na geração do documento: {str(e)}", template_path=self.template_path) from e
    
    def _prepare_replacements(self, data: ExtractedData) -> Dict[str, str]:
        doc_date = self._format_date(data.document.date)
        
        if self._is_pagamento_terceiro_template():
            from .third_party_payment_template_filler import ThirdPartyPaymentTemplateFiller
            filler = ThirdPartyPaymentTemplateFiller()
            replacements = filler.get_pagamento_terceiro_replacements(data)
        else:
            replacements = self.replacement_manager.get_termo_responsabilidade_replacements(data)
        
        if 'ANO_MODELO_VEICULO' in replacements:
            ano_modelo_value = replacements['ANO_MODELO_VEICULO']
            replacements['ANO_TRACKER'] = ano_modelo_value
            replacements['ANO_TRACKER 1.2 TURBO'] = ano_modelo_value
        
        if data.vehicle and data.vehicle.value:
            replacements.update({'{{VALOR_VEICULO_VENDIDO}}': self._format_currency_value(data.vehicle.value), '{{VALOR_AUTORIZADO_PAGAMENTO}}': self._format_currency_value(data.vehicle.value)})
        
        if data.vehicle:
            brand = self._extract_brand_from_model(data.vehicle.model) if data.vehicle.model else ""
            replacements.update({'MARCA_VEICULO_LEGADO': brand, 'MODELO_VEICULO_LEGADO': data.vehicle.model or "", 'CHASSI_VEICULO_LEGADO': data.vehicle.chassis or "", 'COR_VEICULO_LEGADO': data.vehicle.color or "", 'PLACA_VEICULO_LEGADO': data.vehicle.plate or "", 'ANO_MODELO_VEICULO_LEGADO': data.vehicle.year_model or ""})
            if data.vehicle.color:
                for color in ['BRANCA', 'BRANCO', 'PRATA', 'PRETO', 'AZUL', 'VERMELHO', 'CINZA']:
                    replacements[color] = data.vehicle.color
        
        location = self._format_location(data.document.location) if data.document and data.document.location else ''
        replacements.update({'DATA_LOCAL_DOCUMENTO': f"{location}, {doc_date}", 'LOCAL_FORMATADO': location, 'DATA_FORMATADA': doc_date})
        return {k: v for k, v in replacements.items() if v}
    
    def _generate_alternative_filename(self, original_path: str) -> str:
        from datetime import datetime
        import os
        import re
        
        directory = os.path.dirname(original_path)
        filename = os.path.basename(original_path)
        name, ext = os.path.splitext(filename)
        
        # CKDEV-NOTE: Extract first name from current data if available
        first_name = ""
        if self.current_data and hasattr(self.current_data, 'client') and self.current_data.client:
            client_name = self.current_data.client.name if self.current_data.client.name else ""
            if client_name.strip():
                first_name_raw = client_name.strip().split()[0] if client_name.strip() else ""
                first_name = re.sub(r'[^A-Za-zÀ-ÿ]', '', first_name_raw).upper()
        
        timestamp = datetime.now().strftime("%H%M%S")
        
        if first_name:
            alternative_name = f"{name}_v{timestamp}_{first_name}{ext}"
        else:
            alternative_name = f"{name}_v{timestamp}{ext}"
        
        return os.path.join(directory, alternative_name)
    
    def _is_pagamento_terceiro_template(self) -> bool:
        template_name = os.path.basename(self.template_path).lower()
        return 'pagamento_terceiro' in template_name or 'template1_pagamento_terceiro' in template_name
    
    def _extract_brand_from_model(self, model: str) -> str:
        if not model: return ""
        try:
            return get_brand_lookup().get_fallback_brand(model)
        except Exception:
            words = model.strip().split()
            return words[0] if words else ""
    
    def _format_date(self, date_str: str) -> str:
        return datetime.now().strftime("%d/%m/%Y")
    
    def _format_location(self, location: str) -> str:
        if not location: return "São José dos Campos - SP"
        location_clean = re.sub(r'\s*Estado\s*', ' ', location, flags=re.IGNORECASE).strip()
        location_clean = re.sub(r'\s+', ' ', location_clean).strip()
        if '-SP' in location_clean: location_clean = re.sub(r'\s*-\s*SP\s*', ' - SP', location_clean)
        elif location_clean.endswith('SP'): location_clean = re.sub(r'\s+SP\s*$', ' - SP', location_clean)
        elif 'SP' in location_clean: location_clean = re.sub(r'\s+SP\s*', ' - SP', location_clean)
        else: location_clean += ' - SP'
        return re.sub(r'\s+', ' ', location_clean).strip()
    
    
    def _replace_text_in_document(self, doc: Document, replacements: Dict[str, str]):
        # CKDEV-NOTE: Simplified replacement using Termo de Responsabilidade mechanism for all templates
        for paragraph in doc.paragraphs: self._replace_text_in_paragraph(paragraph, replacements)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs: self._replace_text_in_paragraph(paragraph, replacements)
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs: self._replace_text_in_paragraph(paragraph, replacements)
            if section.footer:
                for paragraph in section.footer.paragraphs: self._replace_text_in_paragraph(paragraph, replacements)
    
    def _replace_text_in_paragraph(self, paragraph, replacements: Dict[str, str]):
        # CKDEV-NOTE: Fixed algorithm - only process {{PLACEHOLDER}} style replacements with bold
        if not paragraph.text:
            return
        
        original_text = paragraph.text
        
        # Filter only {{PLACEHOLDER}} style replacements
        placeholder_replacements = {k: v for k, v in replacements.items() 
                                   if k and k.startswith('{{') and k.endswith('}}')}
        
        # Check if any replacements are needed
        has_replacements = any(placeholder in original_text for placeholder in placeholder_replacements.keys())
        if not has_replacements:
            return
        
        # Clear existing runs
        while paragraph.runs:
            paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
        
        # Process text sequentially, splitting at each placeholder
        remaining_text = original_text
        
        while remaining_text:
            # Find the next placeholder in the text
            next_placeholder = None
            next_pos = len(remaining_text)
            
            for placeholder, value in placeholder_replacements.items():
                pos = remaining_text.find(placeholder)
                if pos != -1 and pos < next_pos:
                    next_pos = pos
                    next_placeholder = placeholder
            
            if next_placeholder:
                # Add text before placeholder (normal formatting)
                if next_pos > 0:
                    before_text = remaining_text[:next_pos]
                    run = paragraph.add_run(before_text)
                    run.font.name = "Aptos"
                    run.font.size = Pt(11)
                
                # Add replacement value (bold formatting)
                replacement_value = placeholder_replacements[next_placeholder]
                if replacement_value:
                    bold_run = paragraph.add_run(replacement_value)
                    bold_run.font.name = "Aptos"
                    bold_run.font.size = Pt(11)
                    bold_run.bold = True
                
                # Move to text after placeholder
                remaining_text = remaining_text[next_pos + len(next_placeholder):]
            else:
                # No more placeholders, add remaining text
                if remaining_text:
                    run = paragraph.add_run(remaining_text)
                    run.font.name = "Aptos"
                    run.font.size = Pt(11)
                break
    
    def _replace_text_preserving_format(self, paragraph, old_text: str, new_text: str):
        if old_text not in paragraph.text: return
        full_text = "".join(run.text for run in paragraph.runs)
        if old_text not in full_text: return
        
        for _ in range(len(paragraph.runs)): paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
        parts = full_text.split(old_text, 1)
        
        if len(parts) == 2:
            before_text, after_text = parts
            if before_text:
                before_run = paragraph.add_run(before_text)
                before_run.font.name, before_run.font.size = "Aptos", Pt(11)
            bold_run = paragraph.add_run(new_text)
            bold_run.bold, bold_run.font.name, bold_run.font.size = True, "Aptos", Pt(11)
            if after_text:
                after_run = paragraph.add_run(after_text)
                after_run.font.name, after_run.font.size = "Aptos", Pt(11)
        else:
            run = paragraph.add_run(full_text.replace(old_text, new_text))
            run.bold, run.font.name, run.font.size = True, "Aptos", Pt(11)
    
    
    def _format_currency_value(self, value_str: str) -> str:
        if not value_str or not value_str.strip(): return "R$ 0,00"
        try:
            clean_value = re.sub(r'[^\d.,]', '', value_str.strip())
            if ',' in clean_value and clean_value.count(',') == 1:
                parts = clean_value.split(',')
                if len(parts) == 2 and len(parts[1]) == 2: return f"R$ {clean_value}"
            if '.' in clean_value:
                float_value = float(clean_value)
                formatted = f"{float_value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                return f"R$ {formatted}"
            else:
                num_value = int(clean_value)
                return f"R$ {num_value},00" if num_value <= 999 else f"R$ {num_value:,}".replace(",", ".") + ",00"
        except (ValueError, TypeError):
            return f"R$ {value_str}"