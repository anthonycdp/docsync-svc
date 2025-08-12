import re
from docx import Document
from docx.shared import Pt

def ensure_bold_formatting_for_replacements(doc: Document, replacements: dict):
    for paragraph in doc.paragraphs:
        _ensure_bold_formatting_in_paragraph(paragraph, replacements)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _ensure_bold_formatting_in_paragraph(paragraph, replacements)
    
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                _ensure_bold_formatting_in_paragraph(paragraph, replacements)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                _ensure_bold_formatting_in_paragraph(paragraph, replacements)

def _ensure_bold_formatting_in_paragraph(paragraph, replacements: dict):
    if not paragraph.text:
        return
    
    original_runs_info = []
    for run in paragraph.runs:
        original_runs_info.append({
            'text': run.text,
            'bold': run.bold,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'italic': run.italic,
            'underline': run.underline
        })
    
    has_placeholders = any(placeholder in paragraph.text for placeholder in replacements.keys())
    if not has_placeholders:
        return
    
    while paragraph.runs:
        paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
    
    current_text = "".join(run['text'] for run in original_runs_info)
    
    for placeholder, replacement_value in replacements.items():
        if placeholder in current_text:
            parts = current_text.split(placeholder)
            current_text = ""
            
            for i, part in enumerate(parts):
                if part:
                    run = paragraph.add_run(part)
                    _apply_original_formatting(run, original_runs_info)
                
                if i < len(parts) - 1:
                    bold_run = paragraph.add_run(replacement_value)
                    bold_run.bold = True
                    _apply_original_formatting(bold_run, original_runs_info, force_bold=True)

def _apply_original_formatting(run, original_runs_info, force_bold=False):
    if not original_runs_info:
        run.font.name = "Aptos"
        run.font.size = Pt(11)
        if force_bold:
            run.bold = True
        return
    
    original = original_runs_info[0]
    
    if original['font_name']:
        run.font.name = original['font_name']
    else:
        run.font.name = "Aptos"
    
    if original['font_size']:
        run.font.size = original['font_size']
    else:
        run.font.size = Pt(11)
    
    if force_bold:
        run.bold = True
    elif original['bold'] is not None:
        run.bold = original['bold']
    
    if original['italic'] is not None:
        run.italic = original['italic']
    
    if original['underline'] is not None:
        run.underline = original['underline']

def apply_bold_formatting_to_replaced_values(doc: Document, replacements: dict):
    for paragraph in doc.paragraphs:
        _apply_bold_to_replaced_values_in_paragraph(paragraph, replacements)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _apply_bold_to_replaced_values_in_paragraph(paragraph, replacements)
    
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                _apply_bold_to_replaced_values_in_paragraph(paragraph, replacements)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                _apply_bold_to_replaced_values_in_paragraph(paragraph, replacements)

def _apply_bold_to_replaced_values_in_paragraph(paragraph, replacements: dict):
    if not paragraph.text:
        return
    
    replacement_values = []
    for key, value in replacements.items():
        if value and str(value).strip() and len(str(value).strip()) > 1:
            value_str = str(value).strip()
            if value_str not in ['-', 'SP', 'SÃO PAULO', '0', '00']:
                replacement_values.append(value_str)
    
    if not replacement_values:
        return
    
    has_replaced_values = False
    for replacement_value in replacement_values:
        if replacement_value in paragraph.text:
            has_replaced_values = True
            break
    
    if not has_replaced_values:
        return
    
    current_runs = []
    for run in paragraph.runs:
        current_runs.append({
            'text': run.text,
            'bold': run.bold,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'italic': run.italic,
            'underline': run.underline
        })
    
    while paragraph.runs:
        paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
    
    full_text = "".join(run['text'] for run in current_runs)
    
    replacement_positions = []
    sorted_values = sorted(replacement_values, key=len, reverse=True)
    
    for replacement_value in sorted_values:
        start = 0
        while True:
            pos = full_text.find(replacement_value, start)
            if pos == -1:
                break
            
            overlaps = False
            for existing_start, existing_end, _ in replacement_positions:
                if not (pos >= existing_end or pos + len(replacement_value) <= existing_start):
                    overlaps = True
                    break
            
            if not overlaps:
                replacement_positions.append((pos, pos + len(replacement_value), replacement_value))
            
            start = pos + 1
    
    replacement_positions.sort(key=lambda x: x[0])
    
    if not replacement_positions:
        for run_info in current_runs:
            run = paragraph.add_run(run_info['text'])
            _apply_run_formatting(run, [run_info])
        return
    
    current_pos = 0
    for start, end, replacement_value in replacement_positions:
        if start > current_pos:
            before_text = full_text[current_pos:start]
            if before_text:
                run = paragraph.add_run(before_text)
                _apply_run_formatting(run, current_runs)
        
        bold_run = paragraph.add_run(replacement_value)
        bold_run.bold = True
        _apply_run_formatting(bold_run, current_runs, force_bold=True)
        
        current_pos = end
    
    if current_pos < len(full_text):
        remaining_text = full_text[current_pos:]
        if remaining_text:
            run = paragraph.add_run(remaining_text)
            _apply_run_formatting(run, current_runs)

def _apply_run_formatting(run, reference_runs, force_bold=False):
    if not reference_runs:
        run.font.name = "Aptos"
        run.font.size = Pt(11)
        if force_bold:
            run.bold = True
        return
    
    reference = reference_runs[0]
    
    if reference['font_name']:
        run.font.name = reference['font_name']
    else:
        run.font.name = "Aptos"
    
    if reference['font_size']:
        run.font.size = reference['font_size']
    else:
        run.font.size = Pt(11)
    
    if force_bold:
        run.bold = True
    elif reference['bold'] is not None:
        run.bold = reference['bold']
    
    if reference['italic'] is not None:
        run.italic = reference['italic']
    
    if reference['underline'] is not None:
        run.underline = reference['underline']

def format_placeholders_and_uppercase_bold(docx_path, output_path=None):
    if output_path is None:
        output_path = docx_path
    
    doc = Document(docx_path)
    
    placeholder_pattern = r'\{\{[^}]+\}\}'
    uppercase_pattern = r'\b[A-ZÀÁÂÃÄÅÆÇÈÉÊËÌÍÎÏÐÑÒÓÔÕÖØÙÚÛÜÝÞŸ]{2,}\b'
    
    combined_pattern = f'({placeholder_pattern}|{uppercase_pattern})'
    
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if re.search(combined_pattern, text):
            paragraph.clear()
            
            matches = list(re.finditer(combined_pattern, text))
            
            if matches:
                current_pos = 0
                for match in matches:
                    if match.start() > current_pos:
                        run = paragraph.add_run(text[current_pos:match.start()])
                    
                    match_run = paragraph.add_run(match.group())
                    match_run.bold = True
                    
                    current_pos = match.end()
                
                if current_pos < len(text):
                    run = paragraph.add_run(text[current_pos:])
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if re.search(combined_pattern, text):
                        paragraph.clear()
                        
                        matches = list(re.finditer(combined_pattern, text))
                        
                        if matches:
                            current_pos = 0
                            for match in matches:
                                if match.start() > current_pos:
                                    run = paragraph.add_run(text[current_pos:match.start()])
                                
                                match_run = paragraph.add_run(match.group())
                                match_run.bold = True
                                
                                current_pos = match.end()
                            
                            if current_pos < len(text):
                                run = paragraph.add_run(text[current_pos:])
    
    doc.save(output_path)

def format_placeholders_bold(docx_path, output_path=None):
    if output_path is None:
        output_path = docx_path
    
    doc = Document(docx_path)
    
    placeholder_pattern = r'\{\{[^}]+\}\}'
    
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if re.search(placeholder_pattern, text):
            paragraph.clear()
            
            matches = list(re.finditer(placeholder_pattern, text))
            
            if matches:
                current_pos = 0
                for match in matches:
                    if match.start() > current_pos:
                        run = paragraph.add_run(text[current_pos:match.start()])
                    
                    placeholder_run = paragraph.add_run(match.group())
                    placeholder_run.bold = True
                    
                    current_pos = match.end()
                
                if current_pos < len(text):
                    run = paragraph.add_run(text[current_pos:])
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if re.search(placeholder_pattern, text):
                        paragraph.clear()
                        
                        matches = list(re.finditer(placeholder_pattern, text))
                        
                        if matches:
                            current_pos = 0
                            for match in matches:
                                if match.start() > current_pos:
                                    run = paragraph.add_run(text[current_pos:match.start()])
                                
                                placeholder_run = paragraph.add_run(match.group())
                                placeholder_run.bold = True
                                
                                current_pos = match.end()
                            
                            if current_pos < len(text):
                                run = paragraph.add_run(text[current_pos:])
    
    doc.save(output_path)

def list_placeholders_and_uppercase(docx_path):
    doc = Document(docx_path)
    placeholders = set()
    uppercase_words = set()
    placeholder_pattern = r'\{\{[^}]+\}\}'
    uppercase_pattern = r'\b[A-ZÀÁÂÃÄÅÆÇÈÉÊËÌÍÎÏÐÑÒÓÔÕÖØÙÚÛÜÝÞŸ]{2,}\b'
    
    for paragraph in doc.paragraphs:
        placeholder_matches = re.findall(placeholder_pattern, paragraph.text)
        uppercase_matches = re.findall(uppercase_pattern, paragraph.text)
        placeholders.update(placeholder_matches)
        uppercase_words.update(uppercase_matches)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    placeholder_matches = re.findall(placeholder_pattern, paragraph.text)
                    uppercase_matches = re.findall(uppercase_pattern, paragraph.text)
                    placeholders.update(placeholder_matches)
                    uppercase_words.update(uppercase_matches)
    
    return {
        'placeholders': sorted(list(placeholders)),
        'uppercase_words': sorted(list(uppercase_words))
    }

def list_placeholders(docx_path):
    doc = Document(docx_path)
    placeholders = set()
    placeholder_pattern = r'\{\{[^}]+\}\}'
    
    for paragraph in doc.paragraphs:
        matches = re.findall(placeholder_pattern, paragraph.text)
        placeholders.update(matches)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = re.findall(placeholder_pattern, paragraph.text)
                    placeholders.update(matches)
    
    return sorted(list(placeholders))

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    placeholders = list_placeholders(input_file)
    for placeholder in placeholders:
        pass
    
    format_placeholders_bold(input_file, output_file) 