import os
import re
from typing import Dict
from datetime import datetime
from docx import Document

try:
    from ..data import ExtractedData
    from .template_types import TemplateType, TemplateManager
    from ..utils import get_brand_lookup, LoggerMixin
    from .template_replacements import TemplateReplacementManager
except ImportError:
    import sys
    from pathlib import Path
    sys.path.insert(0, str(Path(__file__).parent.parent))
    from data import ExtractedData
    from processors.template_types import TemplateType, TemplateManager
    from utils import get_brand_lookup, LoggerMixin
    from processors.template_replacements import TemplateReplacementManager

class MultiTemplateProcessor(LoggerMixin):
    
    def __init__(self, templates_dir: str = "templates"):
        super().__init__()
        self.template_manager = TemplateManager(templates_dir)
        self.replacement_manager = TemplateReplacementManager()
        self.current_template_type = None
        self.current_data = None  # CKDEV-NOTE: Store current extracted data for filename generation
    
    def generate_document(self, template_type: TemplateType, data: ExtractedData, output_path: str) -> str:
        self.current_template_type = template_type
        # CKDEV-NOTE: Store current data for alternative filename generation
        self.current_data = data
        config = self.template_manager.get_template_config(template_type)
        if not os.path.exists(config.file_path): 
            raise FileNotFoundError(f"Template não encontrado: {config.file_path}")
        
        doc = None
        try:
            doc = Document(config.file_path)
            replacements = self._prepare_template_replacements(template_type, data)
            self._replace_text_in_document(doc, replacements)
            doc.save(output_path)
            
            try:
                from ..utils.pdf_converter import convert_docx_to_pdf
            except ImportError:
                from utils.pdf_converter import convert_docx_to_pdf
            
            pdf_success, pdf_message, pdf_path = convert_docx_to_pdf(output_path)
            if pdf_success:
                self.logger.info(f"PDF generated successfully: {pdf_path}")
            else:
                self.logger.warning(f"PDF generation failed: {pdf_message}")
            
            return output_path
            
        except PermissionError as e:
            alternative_path = self._generate_alternative_filename(output_path)
            try:
                if doc:
                    doc.save(alternative_path)
                else:
                    doc = Document(config.file_path)
                    replacements = self._prepare_template_replacements(template_type, data)
                    self._replace_text_in_document(doc, replacements)
                    doc.save(alternative_path)
                
                try:
                    from ..utils import convert_docx_to_pdf_libreoffice
                except ImportError:
                    from utils import convert_docx_to_pdf_libreoffice
                
                pdf_success, pdf_message, pdf_path = convert_docx_to_pdf_libreoffice(alternative_path)
                if pdf_success:
                    self.logger.info(f"PDF generated successfully: {pdf_path}")
                else:
                    self.logger.warning(f"PDF generation failed: {pdf_message}")
                
                return alternative_path
            except Exception:
                raise Exception(f"Arquivo bloqueado (provavelmente aberto no Word): {output_path}. Feche o arquivo e tente novamente.") from e
    
    def _prepare_template_replacements(self, template_type: TemplateType, data: ExtractedData) -> Dict[str, str]:
        template_handlers = {
            TemplateType.RESPONSABILIDADE_VEICULO: self._prepare_termo_responsabilidade_replacements,
            TemplateType.PAGAMENTO_TERCEIRO: self._prepare_declaracao_pagamento_replacements,
            TemplateType.CESSAO_CREDITO: self._prepare_termo_dacao_credito_replacements
        }
        
        # CKDEV-NOTE: Search by value to handle different enum instances from serialization
        handler = None
        for key, func in template_handlers.items():
            if key.value == template_type.value:
                handler = func
                break
        
        if handler is None:
            raise ValueError(f"Tipo de template não suportado: {template_type}")
        
        return handler(data)
    
    def _prepare_termo_responsabilidade_replacements(self, data: ExtractedData) -> Dict[str, str]:
        doc_date = self._format_date(data.document.date)
        replacements = self.replacement_manager.get_termo_responsabilidade_replacements(data)
        replacements.update(self._get_legacy_vehicle_replacements(data))
        replacements.update(self._get_location_date_replacements(data, doc_date))
        replacements.update(self._get_vehicle_text_patterns(data))
        return {k: v for k, v in replacements.items() if v}
    
    def _prepare_declaracao_pagamento_replacements(self, data: ExtractedData) -> Dict[str, str]:
        replacements = self.replacement_manager.get_declaracao_pagamento_replacements(data)
        
        if data.payment:
            replacements.update({
                'VALOR_PAGAMENTO_FORMATADO': self._format_payment_amount(data.payment.amount),
                'BANCO_PAGAMENTO': data.payment.bank_name or os.getenv('DEFAULT_BANK_NAME', 'Banco'),
                'CONTA_PAGAMENTO': f"conta corrente {data.payment.account}" if data.payment.account else os.getenv('DEFAULT_ACCOUNT_TYPE', 'conta corrente'),
                'AGENCIA_PAGAMENTO': f"agência {data.payment.agency}" if data.payment.agency else os.getenv('DEFAULT_AGENCY_TYPE', 'agência')
            })
        
        if data.new_vehicle:
            replacements.update({
                'MODELO_VEICULO_NOVO': data.new_vehicle.model or os.getenv('DEFAULT_VEHICLE_MODEL', 'MODELO'),
                'MARCA_VEICULO_NOVO': data.new_vehicle.brand or os.getenv('DEFAULT_VEHICLE_BRAND', 'MARCA'),
                'COR_VEICULO_NOVO': data.new_vehicle.color or os.getenv('DEFAULT_VEHICLE_COLOR', 'COR'),
                'ANO_MODELO_VEICULO_NOVO': data.new_vehicle.year_model or '',
                'PEDIDO_VEICULO_NOVO': data.new_vehicle.sales_order or os.getenv('DEFAULT_SALES_ORDER', 'PEDIDO')
            })
        
        doc_date = self._format_date(data.document.date)
        location = self._format_location(data.document.location) if data.document and data.document.location else ''
        
        replacements.update({
            'NOME_CONCESSIONARIA': os.getenv('DEFAULT_DEALERSHIP_NAME', 'CONCESSIONARIA GENERICA LTDA'),
            'CNPJ_CONCESSIONARIA': os.getenv('DEFAULT_DEALERSHIP_CNPJ', '00.000.000/0000-00'),
            'DATA_DOCUMENTO_FORMATADA': doc_date,
            'LOCAL_DOCUMENTO_FORMATADO': location,
            **self._get_vehicle_text_patterns(data)
        })
        return {k: v for k, v in replacements.items() if v}
    
    def _prepare_termo_dacao_credito_replacements(self, data: ExtractedData) -> Dict[str, str]:
        replacements = self.replacement_manager.get_termo_dacao_credito_replacements(data)
        
        if data.vehicle:
            # CKDEV-NOTE: Consistent brand extraction logic aligned with pp_extractor.py
            brand = data.vehicle.brand
            if not brand or not brand.strip():
                brand = self._extract_brand_from_model(data.vehicle.model) if data.vehicle.model else ""
            replacements.update({
                'MARCA_VEICULO_VENDIDO': brand,
                'MODELO_VEICULO_VENDIDO': data.vehicle.model or '',
                'CHASSI_VEICULO_VENDIDO': data.vehicle.chassis or '',
                'COR_VEICULO_VENDIDO': data.vehicle.color.strip() if data.vehicle.color else os.getenv('DEFAULT_VEHICLE_COLOR', ''),
                'PLACA_VEICULO_VENDIDO': data.vehicle.plate or '',
                'ANO_MODELO_VEICULO_VENDIDO': data.vehicle.year_model or '',
                'VALOR_VEICULO_VENDIDO': f"R$ {data.vehicle.value}" if data.vehicle.value else "R$ 0,00"
            })
        
        if data.new_vehicle:
            replacements.update({
                'MODELO_VEICULO_NOVO': data.new_vehicle.model or os.getenv('DEFAULT_NEW_VEHICLE_MODEL', 'MODELO NOVO'),
                'MARCA_VEICULO_NOVO': data.new_vehicle.brand or os.getenv('DEFAULT_NEW_VEHICLE_BRAND', 'MARCA NOVA'),
                'ANO_MODELO_VEICULO_NOVO': data.new_vehicle.year_model or '',
                'CHASSI_VEICULO_NOVO': data.new_vehicle.chassis or os.getenv('DEFAULT_NEW_VEHICLE_CHASSIS', 'CHASSI_NOVO'),
                'VALOR_VEICULO_NOVO': f"R$ {data.new_vehicle.value}" if data.new_vehicle and data.new_vehicle.value else "R$ 0,00"
            })
        
        doc_date = self._format_date(data.document.date)
        location = self._format_location(data.document.location) if data.document and data.document.location else ''
        
        replacements.update({
            'CNPJ_CONCESSIONARIA': os.getenv('DEFAULT_DEALERSHIP_CNPJ', '00.000.000/0000-00'),
            'ENDERECO_CONCESSIONARIA': os.getenv('DEFAULT_DEALERSHIP_ADDRESS', 'ENDERECO CONCESSIONARIA GENERICA'),
            'DADOS_BANCARIOS': os.getenv('DEFAULT_BANK_DATA', 'DADOS_BANCARIOS_GENERICOS'),
            'DATA_DOCUMENTO_DACAO': doc_date,
            'LOCAL_DOCUMENTO_DACAO': location
        })
        
        # CKDEV-NOTE: Add vehicle text patterns with safe filtering for Cessão de Crédito
        vehicle_patterns = self._get_vehicle_text_patterns(data)
        for pattern_key, pattern_value in vehicle_patterns.items():
            # Only add patterns that won't break formatting
            if pattern_key and pattern_value and len(pattern_key) < 200:  # Limit pattern length
                replacements[pattern_key] = pattern_value
        
        return {k: v for k, v in replacements.items() if k and v}
    
    def _format_payment_amount(self, amount: str) -> str:
        if not amount:
            return ""
        try:
            return f"R$ {amount} (Valor por extenso)" if float(amount.replace(',', '.').replace('.', '')) else f"R$ {amount}"
        except:
            return ""
    
    def _get_legacy_vehicle_replacements(self, data: ExtractedData) -> Dict[str, str]:
        if not data.vehicle:
            return {}
        
        # CKDEV-NOTE: Consistent brand extraction logic aligned with pp_extractor.py
        brand = data.vehicle.brand
        if not brand or not brand.strip():
            brand = self._extract_brand_from_model(data.vehicle.model) if data.vehicle.model else ""
        replacements = {
            'MARCA_VEICULO': brand,
            'MODELO_VEICULO': data.vehicle.model or "",
            'CHASSI_VEICULO': data.vehicle.chassis or "",
            'COR_VEICULO': data.vehicle.color.strip() if data.vehicle.color else "",
            'PLACA_VEICULO': data.vehicle.plate or "",
            'ANO_MODELO_VEICULO': data.vehicle.year_model or ""
        }
        
        if data.vehicle.color:
            default_colors = ['BRANCA', 'BRANCO', 'PRATA', 'PRETO', 'AZUL', 'VERMELHO', 'CINZA']
            configurable_colors = os.getenv('VEHICLE_COLORS', '').split(',') if os.getenv('VEHICLE_COLORS') else default_colors
            configurable_colors = [color.strip() for color in configurable_colors if color.strip()]
            
            for color in configurable_colors:
                replacements[color] = data.vehicle.color
        
        return replacements
    
    def _get_location_date_replacements(self, data: ExtractedData, doc_date: str) -> Dict[str, str]:
        location = self._format_location(data.document.location) if data.document and data.document.location else ''
        date_location = f"{location.rstrip()}, {doc_date}"
        date_location = re.sub(r'\s+,', ',', date_location)
        return {
            'LOCALIZACAO_DATA': date_location,
            'CIDADE_DATA': date_location,
            'LOCAL_DOCUMENTO': location,
            'DATA_DOCUMENTO': doc_date
        }

    def _extract_brand_from_model(self, model: str) -> str:
        if not model:
            return ""
        try:
            return get_brand_lookup().get_brand_from_model(model) or ""
        except:
            return ""
    
    def _format_date(self, date_str: str) -> str:
        return datetime.now().strftime("%d/%m/%Y")
    
    def _get_vehicle_text_patterns(self, data: ExtractedData) -> Dict[str, str]:
        if not data.vehicle:
            return {}
            
        # CKDEV-NOTE: Consistent brand extraction logic aligned with pp_extractor.py
        brand = data.vehicle.brand
        if not brand or not brand.strip():
            brand = self._extract_brand_from_model(data.vehicle.model) if data.vehicle.model else ""
        model, chassis, color, plate, year_model = data.vehicle.model or "", data.vehicle.chassis or "", data.vehicle.color or "", data.vehicle.plate or "", data.vehicle.year_model or ""
        
        # CKDEV-NOTE: Preserve formatting consistency - generate clean formatted text
        if hasattr(data.vehicle, 'format_with_commas'):
            formatted_text = data.vehicle.format_with_commas()
        else:
            # Manual formatting with consistent structure
            parts = []
            if brand: parts.append(brand.strip())
            if model: parts.append(f"MODELO {model.strip()}")
            if chassis: parts.append(f"CHASSI {chassis.strip()}")
            if color: parts.append(f"COR {color.strip()}")
            if plate: parts.append(f"PLACA {plate.strip()}")
            if year_model: parts.append(f"ANO/MODELO {year_model.strip()}")
            formatted_text = ", ".join(parts) if parts else ""
        
        # CKDEV-NOTE: Generate only safe patterns that won't break document formatting
        patterns = {}
        
        # Only create patterns if we have sufficient vehicle data and pattern length is reasonable
        if brand and model and chassis and color and plate and year_model:
            full_pattern = f'{brand.strip()} MODELO {model.strip()} CHASSI {chassis.strip()} COR {color.strip()} PLACA {plate.strip()} ANO/MODELO {year_model.strip()}'
            if len(full_pattern) < 150:  # Prevent overly long patterns
                patterns[full_pattern] = formatted_text
        
        if model and chassis and color and plate and year_model:
            partial_pattern = f'MODELO {model.strip()} CHASSI {chassis.strip()} COR {color.strip()} PLACA {plate.strip()} ANO/MODELO {year_model.strip()}'
            if len(partial_pattern) < 150:  # Prevent overly long patterns
                patterns[partial_pattern] = formatted_text
        
        # CKDEV-NOTE: Filter out problematic patterns that could break formatting
        return {k: v for k, v in patterns.items() if k and v and len(k.strip()) > 0 and len(k) < 200}
    
    def _format_location(self, location: str) -> str:
        if not location:
            return ''
        
        location_clean = re.sub(r'\s*Estado\s*', ' ', location, flags=re.IGNORECASE).strip()
        location_clean = re.sub(r'\s+', ' ', location_clean).strip()
        
        if '-SP' in location_clean:
            location_clean = re.sub(r'\s*-\s*SP\s*', ' - SP', location_clean)
        elif location_clean.endswith('SP'):
            location_clean = re.sub(r'\s+SP\s*$', ' - SP', location_clean)
        elif 'SP' in location_clean:
            location_clean = re.sub(r'\s+SP\s*', ' - SP', location_clean)
        else:
            location_clean += ' - SP'
        
        return re.sub(r'\s+', ' ', location_clean).strip()
    
    def _generate_alternative_filename(self, original_path: str) -> str:
        from datetime import datetime
        import re
        
        directory = os.path.dirname(original_path)
        filename = os.path.basename(original_path)
        name, ext = os.path.splitext(filename)
        
        # CKDEV-NOTE: Extract first name from current data if available
        first_name = ""
        if self.current_data and hasattr(self.current_data, 'client') and self.current_data.client:
            client_name = self.current_data.client.name if self.current_data.client.name else ""
            if client_name.strip():
                first_name_raw = client_name.strip().split()[0] if client_name.strip() else ""
                first_name = re.sub(r'[^A-Za-zÀ-ÿ]', '', first_name_raw).upper()
        
        timestamp = datetime.now().strftime("%H%M%S")
        
        if first_name:
            alternative_name = f"{name}_v{timestamp}_{first_name}{ext}"
        else:
            alternative_name = f"{name}_v{timestamp}{ext}"
        
        return os.path.join(directory, alternative_name)
    
    
    
    def _replace_text_in_document(self, doc: Document, replacements: Dict[str, str]):
        # CKDEV-NOTE: Consistent text replacement mechanism for all templates with formatting preservation
        # Process body paragraphs
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, replacements)
        
        # Process table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_text_in_paragraph(paragraph, replacements)
        
        # Process headers and footers
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    self._replace_text_in_paragraph(paragraph, replacements)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    self._replace_text_in_paragraph(paragraph, replacements)
    
    def _replace_text_in_paragraph(self, paragraph, replacements: Dict[str, str]):
        # CKDEV-NOTE: Preserve original formatting while replacing placeholders
        if not paragraph.text:
            return
        
        original_text = paragraph.text
        
        # Filter only {{PLACEHOLDER}} style replacements
        placeholder_replacements = {k: v for k, v in replacements.items() 
                                   if k and k.startswith('{{') and k.endswith('}}')}
        
        # Check if any replacements are needed
        has_replacements = any(placeholder in original_text for placeholder in placeholder_replacements.keys())
        if not has_replacements:
            return
        
        # CKDEV-NOTE: Preserve original run formatting when possible
        # Store original run formatting before clearing
        original_runs = []
        for run in paragraph.runs:
            original_runs.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
            })
        
        # Get default formatting from first run if exists
        default_format = original_runs[0] if original_runs else {
            'bold': False,
            'italic': False,
            'underline': False,
            'font_name': None,
            'font_size': None,
            'font_color': None
        }
        
        # Clear existing runs
        while paragraph.runs:
            paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
        
        # Process text sequentially, splitting at each placeholder
        remaining_text = original_text
        
        while remaining_text:
            # Find the next placeholder in the text
            next_placeholder = None
            next_pos = len(remaining_text)
            
            for placeholder, value in placeholder_replacements.items():
                pos = remaining_text.find(placeholder)
                if pos != -1 and pos < next_pos:
                    next_pos = pos
                    next_placeholder = placeholder
            
            if next_placeholder:
                # Add text before placeholder (preserve original formatting)
                if next_pos > 0:
                    before_text = remaining_text[:next_pos]
                    run = paragraph.add_run(before_text)
                    # Apply original formatting
                    if default_format['font_name']:
                        run.font.name = default_format['font_name']
                    if default_format['font_size']:
                        run.font.size = default_format['font_size']
                    if default_format['italic'] is not None:
                        run.italic = default_format['italic']
                    if default_format['underline'] is not None:
                        run.underline = default_format['underline']
                    if default_format['font_color']:
                        run.font.color.rgb = default_format['font_color']
                
                # Add replacement value (bold formatting but preserve other styles)
                replacement_value = placeholder_replacements[next_placeholder]
                if replacement_value:
                    bold_run = paragraph.add_run(replacement_value)
                    # Apply original formatting plus bold
                    if default_format['font_name']:
                        bold_run.font.name = default_format['font_name']
                    if default_format['font_size']:
                        bold_run.font.size = default_format['font_size']
                    bold_run.bold = True
                    if default_format['italic'] is not None:
                        bold_run.italic = default_format['italic']
                    if default_format['underline'] is not None:
                        bold_run.underline = default_format['underline']
                    if default_format['font_color']:
                        bold_run.font.color.rgb = default_format['font_color']
                
                # Move to text after placeholder
                remaining_text = remaining_text[next_pos + len(next_placeholder):]
            else:
                # No more placeholders, add remaining text with original formatting
                if remaining_text:
                    run = paragraph.add_run(remaining_text)
                    # Apply original formatting
                    if default_format['font_name']:
                        run.font.name = default_format['font_name']
                    if default_format['font_size']:
                        run.font.size = default_format['font_size']
                    if default_format['italic'] is not None:
                        run.italic = default_format['italic']
                    if default_format['underline'] is not None:
                        run.underline = default_format['underline']
                    if default_format['font_color']:
                        run.font.color.rgb = default_format['font_color']
                break
    
    
    def _replace_text_with_bold_formatting(self, paragraph, old_text: str, new_text: str):
        full_text = "".join(run.text for run in paragraph.runs)
        if old_text not in full_text:
            return
        for _ in range(len(paragraph.runs)):
            paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
        parts = full_text.split(old_text, 1)
        if len(parts) == 2:
            before_text, after_text = parts
            if before_text:
                paragraph.add_run(before_text)
            bold_run = paragraph.add_run(new_text)
            bold_run.bold = True
            if after_text:
                paragraph.add_run(after_text)
        else:
            if old_text in full_text:
                start_pos = full_text.find(old_text)
                before, after = full_text[:start_pos], full_text[start_pos + len(old_text):]
                if before:
                    paragraph.add_run(before)
                bold_run = paragraph.add_run(new_text)
                bold_run.bold = True
                if after:
                    paragraph.add_run(after)
            else:
                run = paragraph.add_run(full_text.replace(old_text, new_text))
                run.bold = True
    
    
    def _format_third_party_address(self, third_party) -> str:
        if not third_party:
            return ""
        address_part = third_party.address.upper().strip() if third_party.address else ""
        city_part = third_party.city.upper().strip() if third_party.city else ""
        cep_part = third_party.cep.strip() if third_party.cep else ""
        clean_address = re.sub(r',?\s*CEP\s*\d{5}-?\d{3}.*$', '', address_part).strip() if "CEP" in address_part else address_part
        return (f"{clean_address}, CEP {cep_part}, CIDADE/ESTADO: {city_part}" if clean_address and city_part and cep_part 
                else f"{clean_address}, CIDADE/ESTADO: {city_part}" if clean_address and city_part 
                else clean_address if clean_address else "")
    
    def process_template(self, template_type: TemplateType, data: ExtractedData, output_path: str) -> str:
        return self.generate_document(template_type, data, output_path)